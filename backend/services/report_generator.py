from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os


class SingleReportGenerator:
    def generate(self, results: dict, output_path: str, image_path: str = None):
        doc = Document()

        title = doc.add_heading("排水管道缺陷智能评估报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.add_run(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta.add_run("评估标准：CJJ 181-2012《城镇排水管道检测与评估技术规程》\n")
        meta.add_run("评估方式：AI辅助评估（辅助决策参考，不替代正式检测成果表）\n")

        if image_path and os.path.exists(image_path):
            doc.add_heading("检测图像", level=1)
            doc.add_picture(image_path, width=Inches(5.5))

        self._write_defect_sections(doc, results)
        self._write_disclaimer(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)

    def _write_defect_sections(self, doc, results):
        summary = results.get("summary", {})
        doc.add_heading("一、缺陷概况", level=1)
        doc.add_paragraph(
            f"本次评估共检测到 {summary.get('total_count', 0)} 处缺陷，"
            f"涉及 {', '.join(summary.get('defect_types', []))} 等类型。"
        )

        doc.add_heading("二、逐项评估详情", level=1)
        grade_colors = {
            1: RGBColor(0, 128, 0),
            2: RGBColor(255, 165, 0),
            3: RGBColor(255, 128, 0),
            4: RGBColor(255, 0, 0),
        }

        for i, defect in enumerate(results.get("defects", []), 1):
            doc.add_heading(
                f"缺陷 {i}：{defect.get('name', '')} ({defect.get('code', '')})",
                level=2,
            )
            doc.add_paragraph(f"缺陷描述：{defect.get('description', '')}")
            gp = doc.add_paragraph()
            gr = gp.add_run(f"等级判定：{defect.get('grade', '?')}级")
            gr.bold = True
            gr.font.color.rgb = grade_colors.get(
                defect.get("grade"), RGBColor(0, 0, 0)
            )
            doc.add_paragraph(f"判定依据：{defect.get('criteria', '')}")
            doc.add_paragraph(f"特征分析：{defect.get('analysis', '')}")

        doc.add_heading("三、置信度说明", level=1)
        for d in results.get("defects", []):
            doc.add_paragraph(
                f"{d.get('name', '')}缺陷评估置信度：{d.get('confidence_level', '中')}",
                style="List Bullet",
            )

        doc.add_heading("四、修复建议与优先级", level=1)
        for d in sorted(
            results.get("defects", []), key=lambda x: x.get("grade", 0), reverse=True
        ):
            doc.add_paragraph(
                f"[{d.get('grade', '?')}级] {d.get('name', '')}：{d.get('repair_suggestion', '')}",
                style="List Bullet",
            )

    def _write_disclaimer(self, doc):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("声明：").bold = True
        p.add_run(
            "本报告由AI系统自动生成，定位为面向辅助决策的可解释性评估意见，"
            "不替代CJJ 181-2012标准附录D中定义的正式检测成果表。"
            "建议由持证技术人员审核确认后使用。"
        )


class BatchReportGenerator:
    def generate(self, all_assessments: list, output_path: str):
        doc = Document()

        title = doc.add_heading("排水管道缺陷批量评估报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.add_run(
            f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        )
        meta.add_run(f"评估图像数量：{len(all_assessments)} 张\n")
        meta.add_run("评估标准：CJJ 181-2012《城镇排水管道检测与评估技术规程》\n")
        meta.add_run("评估方式：AI辅助批量评估（辅助决策参考）\n")

        doc.add_heading("总体统计", level=1)
        self._write_summary_table(doc, all_assessments)

        doc.add_heading("逐图评估详情", level=1)
        for i, assessment in enumerate(all_assessments, 1):
            self._write_single_image_section(doc, i, assessment)

        self._write_disclaimer(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)

    def _write_summary_table(self, doc, all_assessments):
        total_defects = 0
        type_dist = {}
        error_count = 0

        for a in all_assessments:
            if a["status"] == "error":
                error_count += 1
                continue
            for d in a.get("detection", {}).get("detections", []):
                total_defects += 1
                code = d["class_code"]
                type_dist[code] = type_dist.get(code, 0) + 1

        doc.add_paragraph(f"总处理图像数：{len(all_assessments)} 张")
        doc.add_paragraph(f"检测到的缺陷总数：{total_defects} 处")
        doc.add_paragraph(f"处理失败数：{error_count} 张")

        if type_dist:
            doc.add_heading("缺陷类型分布", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "缺陷类型", "代码", "数量"
            name_map = {
                "PL": "破裂",
                "CK": "错口",
                "SG": "树根",
                "SL": "渗漏",
                "TL": "脱落",
                "ZW": "障碍物",
            }
            for code, count in sorted(type_dist.items(), key=lambda x: -x[1]):
                row = table.add_row().cells
                row[0].text = name_map.get(code, code)
                row[1].text = code
                row[2].text = str(count)

    def _write_single_image_section(self, doc, index, assessment):
        doc.add_heading(f"图像 {index}：{assessment['image_name']}", level=2)

        img_path = assessment["image_path"]
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(4.5))

        detection = assessment.get("detection", {})
        defects = detection.get("detections", [])
        doc.add_paragraph(f"检测到 {len(defects)} 处缺陷")

        if defects:
            dt = doc.add_table(rows=1, cols=4)
            dt.style = "Table Grid"
            h = dt.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = (
                "缺陷类型",
                "置信度",
                "位置",
                "坐标",
            )
            for d in defects:
                r = dt.add_row().cells
                r[0].text = d["class_name"]
                r[1].text = f"{d['confidence']*100:.1f}%"
                r[2].text = d["position_description"]
                b = d["bbox_pixels"]
                r[3].text = f"({b['x1']},{b['y1']})-({b['x2']},{b['y2']})"

        if assessment["status"] == "success" and assessment.get("assessment_text"):
            doc.add_heading("AI评估结果", level=3)
            for para_text in assessment["assessment_text"].split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        elif assessment["status"] == "error":
            doc.add_paragraph(f"评估失败：{assessment.get('error', '未知错误')}")

        doc.add_page_break()

    def _write_disclaimer(self, doc):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("声明：").bold = True
        p.add_run(
            "本报告由AI系统批量自动生成，定位为面向辅助决策的可解释性评估意见。"
            "本报告不是CJJ 181-2012标准附录D中定义的检测成果表，"
            "不具备工程归档效力。建议由持证技术人员逐条审核确认后使用。"
        )
